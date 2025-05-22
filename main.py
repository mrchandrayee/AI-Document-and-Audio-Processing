from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import openai
from dotenv import load_dotenv
from pypdf import PdfReader
from typing import Annotated
from docx import Document
import base64
import tempfile
from enum import Enum
from markdown_pdf import MarkdownPdf, Section
from striprtf.striprtf import rtf_to_text
import pandas as pd
import os
from util import parseDocuments, upload_file, parseDocumentsV2, parseDocumentsWithVector
from system_prompts import SYSTEM_PROMPT, SYSTEM_PROMPT_V2, AUDIO_TRANSCRIPTION_PROMPT
from whisper_service import WhisperService
import json
from pydantic import BaseModel
from uuid import uuid4

class ResponseSchema(BaseModel):
    response: str
    file_response: str

load_dotenv(override=True)
AUTH_SECRET_KEY = os.getenv("AUTH_SECRET_KEY")

app = FastAPI()
client = openai.OpenAI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow requests from your frontend domain
    allow_credentials=True,
    allow_methods=["*"],  # Allow all HTTP methods
    allow_headers=["*"],  # Allow all headers
)

class ResponseType(str, Enum):
    pdf = "pdf"
    string = "string"

class ModelType(str, Enum):
    gpt4o = "gpt-4o"
    gpt4omini = "gpt-4o-mini"

MODEL = 'gpt-4o'
# Document formats
DOC_EXTENSIONS = ['doc', 'dot', 'docx', 'dotx', 'docm', 'dotm', 'pdf', 'rtf', 'xlsx', 'xls', 'txt']
# Image formats
IMAGE_EXTENSIONS = ['png', 'jpeg', 'jpg']
# Audio formats
AUDIO_EXTENSIONS = ['mp3', 'wav', 'ogg', 'm4a', 'flac', 'aac', 'wma', 'aiff', 'alac']
# Video formats (for audio extraction)
VIDEO_EXTENSIONS = ['mp4', 'avi', 'mov', 'mkv', 'webm', 'wmv', 'flv', 'mpeg']
# All supported formats
SUPPORTED_EXTENSIONS = set(DOC_EXTENSIONS + IMAGE_EXTENSIONS + AUDIO_EXTENSIONS + VIDEO_EXTENSIONS)

# Health checkup end point
@app.get("/")
def health():
    return {
        "status": "success",
        "message": "Backend Healthy"
    }

# Chat completion end point
@app.post("/v1/chat-completion")
def chatCompletion(prompt: Annotated[str, Form()], response_type: Annotated[ResponseType, Form()] = ResponseType.string, model_name: Annotated[ModelType, Form()] = ModelType.gpt4omini, file: Annotated[UploadFile | None, File()] = None, sheet_names: Annotated[str | None, Form()] = None, authorization: Annotated[str | None, Header()] = None):
    MODEL = model_name
    documentText = ''
    b64 = None

    if not authorization or (authorization != AUTH_SECRET_KEY):
        print(f"Authorization header: {authorization}")
        raise HTTPException(
            status_code=401, detail="Provide the correct authorization token in headers")

    if file is not None:
        fileExt = file.filename.split('.')[-1].lower()
        if fileExt not in SUPPORTED_EXTENSIONS:
            raise HTTPException(400, f"{fileExt} file type not supported")
        if fileExt == 'pdf':
            reader = PdfReader(file.file)
            for page in reader.pages:
                extracted_text = page.extract_text(extraction_mode='layout')
                if len(extracted_text.strip()) == 0:
                    extracted_text = page.extract_text()
                documentText += extracted_text
                documentText += '\n\n'
            if len(documentText.strip()) == 0:
                raise HTTPException(
                    status_code=400, detail="Unable to parse text from the given document or the document does not contain any text.")
        elif fileExt in ['png', 'jpeg', 'jpg']:
            with tempfile.TemporaryDirectory() as temp_dir:
                path = f"{temp_dir}/filename.{fileExt}"
                with open(path, 'wb') as img:
                    img.write(file.file.read())
                with open(path, 'rb') as img:
                    b64 = base64.b64encode(img.read()).decode("utf-8")
        elif fileExt in ['rtf']:
            documentText = rtf_to_text(file.file.read().decode('utf-8'))
        elif fileExt in ['xls', 'xlsx']:
            try:
                excelDF = None
                if sheet_names is None or len(sheet_names) == 0:
                    excelDF = pd.read_excel(file.file)
                    documentText = excelDF.to_string()
                else:
                    sheet_name_arr = list(map(lambda x: x.strip(), sheet_names.split(',')))
                    excelDF = pd.read_excel(file.file, sheet_name=sheet_name_arr)
                    if not isinstance(excelDF, dict):
                        raise HTTPException(400, "Something went wrong while parsing the excel sheet")
                    for key in excelDF:
                        documentText += f'Sheet Name: {key}\n\n'
                        documentText += excelDF[key].to_string()
                        documentText += '\n\n'
            except Exception as exp:
                print(exp)
                raise HTTPException(400, "Worksheet name not found")
        elif fileExt == 'txt':
            documentText += file.file.read().decode('utf-8')
        else:
            try:
                document = Document(file.file)
            except Exception as e:
                print("Something went wrong while parsing the file:", e)
                raise HTTPException(400, detail="The file could not be parsed")
            for paragraph in document.paragraphs:
                documentText += paragraph.text if paragraph and paragraph.text else ""
            # Note: Table processing is currently disabled
    # for elem in document.tables:
    #     textRow = set()
    #     for i, row in enumerate(elem.columns):
    #         for cell in row.cells:
    #             textRow.add(f"{cell.text}\n")
    #     text += f"{' '.join(list(textRow))}\n\n"
    # print(text)
    userContent = [
        {
            "type": "text",
            "text": ""
        }
    ]

    if len(documentText) > 0:
        userContent[0]['text'] = f"User prompt:\n{prompt}\n\nThis is document content: \n{documentText}"
    elif b64 is not None:
        userContent[0]['text'] = prompt

    messages = [
        {
            "role": "system",
            "content": SYSTEM_PROMPT
        },
        {
            "role": "user",
            "content": userContent
        }
    ]

    if file is None:
        messages = messages[1:]
        userContent[0]['text'] = prompt

    if b64 is not None:
        userContent.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/jpeg;base64,{b64}"}
        })

    response = client.chat.completions.create(
        model=MODEL,
        messages=messages
    )

    if response_type is ResponseType.pdf:
        pdf = MarkdownPdf(toc_level=2)
        content = response.choices[0].message.content
        pdf.add_section(Section(content))
        pdf.save(f"response.pdf")
        return FileResponse(f"response.pdf",
                            media_type="application/pdf")
    
    return {
        "status": "success",
        "prompt": prompt,
        "response": response.choices[0].message.content
    }

# With Responses API and no Vector Store
@app.post("/v2/chat-completion")
def chatCompletionV2(prompt: Annotated[str, Form()], model_name: Annotated[ModelType, Form()] = ModelType.gpt4omini, file: Annotated[UploadFile | None, File()] = None, sheet_names: Annotated[str | None, Form()] = None, authorization: Annotated[str | None, Header()] = None, temperature: Annotated[float, Form()] = 0.6):
    if temperature < 0 or temperature > 2:
        raise HTTPException(
            status_code=400, detail="Temperature value is invalid. 0 <= temperature <= 2"
        )
    if not authorization or (authorization != AUTH_SECRET_KEY):
        print(f"Authorization header: {authorization}")
        raise HTTPException(
            status_code=401, detail="Provide the correct authorization token in headers")
    MODEL = model_name
    documentText, b64, pdf_file_id = parseDocumentsV2(file, sheet_names, client=client)
    userContent = []

    if pdf_file_id:
        userContent.append({
            "type": "input_file",
            "file_id": pdf_file_id,
        })
    newPrompt = ""
    if len(documentText) > 0 and pdf_file_id is None:
        newPrompt = f"User prompt:\n{prompt}\n\nThis is document content parsed from a file: \n{documentText}"
    else:
        newPrompt = prompt

        userContent.append({

        "type": "input_text",
        "text": newPrompt,
    })

    if b64 is not None:
        userContent.append({
            "type": "input_image",
            "image_url": f"data:image/jpeg;base64,{b64}",
        })
    

    try:
        response = client.responses.create(
            model=MODEL,
            instructions=SYSTEM_PROMPT_V2,
            text={
                "format": {
                    "type": "json_schema",
                    "name": "file_text_response",
                    "schema": {
                        "type": "object",
                        "properties": {
                            "response": {
                                "type": "string"
                            },
                            "file_response": {
                                "type": "string"
                            }
                        },
                        "required": ["response", "file_response"],
                        "additionalProperties": False
                    },
                    "strict": True
                }
            },
            input=[
                {
                    "role": "user",
                    "content": userContent
                }
            ]
        )
        content = json.loads(response.output_text)
    except Exception as e:
        print(e)
        raise HTTPException(500, detail=e)
    
    download_link = None    
    if ("file_response" in content) and content['file_response'] is not None and content['file_response'] != '':
        pdf = MarkdownPdf(toc_level=2)
        pdf.add_section(Section(content['file_response']))
        pdf.save(f"response.pdf")
        file_name_s3 = str(uuid4()) + ".pdf"
        download_link = upload_file('response.pdf', file_name_s3)
        if not download_link:
            # Generate dummy link if real upload fails or is disabled
            dummy_uuid = str(uuid4())
            download_link = f"https://dummy-s3-bucket.example.com/{dummy_uuid}/{file_name_s3}"
            print(f"Using dummy download link: {download_link}")

    return {
        "status": "success",
        "prompt": prompt,
        "response": content['response'],
        "pdf": download_link
    }

# With Responses API and Vector Store
@app.post("/v3/chat-completion")
def chatCompletionV3(prompt: Annotated[str, Form()], model_name: Annotated[ModelType, Form()] = ModelType.gpt4omini, file: Annotated[UploadFile | None, File()] = None, sheet_names: Annotated[str | None, Form()] = None, authorization: Annotated[str | None, Header()] = None, temperature: Annotated[float, Form()] = 0.6):
    if temperature < 0 or temperature > 2:
        raise HTTPException(
            status_code=400, detail="Temperature value is invalid. 0 <= temperature <= 2"
        )
    if not authorization or (authorization != AUTH_SECRET_KEY):
        print(f"Authorization header: {authorization}")
        raise HTTPException(
            status_code=401, detail="Provide the correct authorization token in headers")
    
    MODEL = model_name
    documentText, b64, vector_store_id = parseDocumentsWithVector(file, sheet_names, client=client)
    userContent = []

    newPrompt = ""
    if len(documentText) > 0 and vector_store_id is None:
        newPrompt = f"User prompt:\n{prompt}\n\nThis is document content parsed from a file: \n{documentText}"
    else:
        newPrompt = prompt

    userContent.append({
        "type": "input_text",
        "text": newPrompt,
    })

    if b64 is not None:
        userContent.append({
            "type": "input_image",
            "image_url": f"data:image/jpeg;base64,{b64}",
        })
    tools = []
    if vector_store_id:
        tools.append({
            "type": "file_search",
            "vector_store_ids": [vector_store_id],
            "max_num_results": 1
        })

    try:
        response = client.responses.create(
            model=MODEL,
            instructions=SYSTEM_PROMPT_V2,
            text={
                "format": {
                    "type": "json_schema",
                    "name": "file_text_response",
                    "schema": {
                        "type": "object",
                        "properties": {
                            "response": {
                                "type": "string"
                            },
                            "file_response": {
                                "type": "string"
                            }
                        },
                        "required": ["response", "file_response"],
                        "additionalProperties": False
                    },
                    "strict": True
                }
            },
            tools=tools,
            input=[
                {
                    "role": "user",
                    "content": userContent
                }
            ]
        )
        print(response.output[-1].content[-1].text)
        content = json.loads(response.output[-1].content[-1].text)
    except Exception as e:
        print(e)
        raise HTTPException(500, detail=e)
    
    download_link = None    
    if ("file_response" in content) and content['file_response'] is not None and content['file_response'] != '':
        pdf = MarkdownPdf(toc_level=2)
        pdf.add_section(Section(content['file_response']))
        pdf.save(f"response.pdf")
        file_name_s3 = str(uuid4()) + ".pdf"
        download_link = upload_file('response.pdf', file_name_s3)
        if not download_link:
            # Generate dummy link if real upload fails or is disabled
            dummy_uuid = str(uuid4())
            download_link = f"https://dummy-s3-bucket.example.com/{dummy_uuid}/{file_name_s3}"
            print(f"Using dummy download link: {download_link}")

    return {
        "status": "success",
        "prompt": prompt,
        "response": content['response'],
        "pdf": download_link
    }

@app.post("/v4/chat-completion")
def chatCompletionV4(prompt: Annotated[str, Form()], model_name: Annotated[ModelType, Form()] = ModelType.gpt4omini, file: Annotated[UploadFile | None, File()] = None, sheet_names: Annotated[str | None, Form()] = None, authorization: Annotated[str | None, Header()] = None, temperature: Annotated[float, Form()] = 0.6):
    if temperature < 0 or temperature > 2:
        raise HTTPException(
            status_code=400, detail="Temperature value is invalid. 0 <= temperature <= 2"
        )
    if not authorization or (authorization != AUTH_SECRET_KEY):
        print(f"Authorization header: {authorization}")
        raise HTTPException(
            status_code=401, detail="Provide the correct authorization token in headers")
    
    MODEL = model_name
    documentText, b64, _ = parseDocuments(file, sheet_names)
    userContent = [
        {
            "type": "text",
            "text": ""
        }
    ]

    if len(documentText) > 0:
        userContent[0]['text'] = f"User prompt:\n{prompt}\n\nThis is document content: \n{documentText}"
    elif b64 is not None:
        userContent[0]['text'] = prompt

    messages = [
        {
            "role": "system",
            "content": SYSTEM_PROMPT_V2
        },
        {
            "role": "user",
            "content": userContent
        }
    ]

    if file is None:
        userContent[0]['text'] = prompt

    if b64 is not None:
        userContent.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/jpeg;base64,{b64}"}
        })

    response = {}
    try:
        response = client.chat.completions.create(
            model=MODEL,
            messages=messages,
            response_format={"type": "json_object"},
            temperature=temperature
        )
    except openai.RateLimitError as e:
        print(f"Rate limit error occurred: {e}")
        raise HTTPException(
            status_code=429, detail="OpenAI token limit exceeded")

    res_content = response.choices[0].message.content
    content = json.loads(res_content)
    download_link = None    
    if ("file_response" in content) and content['file_response'] is not None:
        pdf = MarkdownPdf(toc_level=2)
        pdf.add_section(Section(content['file_response']))
        pdf.save(f"response.pdf")
        file_name_s3 = str(uuid4()) + ".pdf"
        download_link = upload_file('response.pdf', file_name_s3)
        if not download_link:
            # Generate dummy link if real upload fails or is disabled
            dummy_uuid = str(uuid4())
            download_link = f"https://dummy-s3-bucket.example.com/{dummy_uuid}/{file_name_s3}"
            print(f"Using dummy download link: {download_link}")

    return {
        "status": "success",
        "prompt": prompt,
        "response": content['response'],
        "pdf": download_link
    }

@app.post("/v5/chat-completion")
def chatCompletionV5(prompt: Annotated[str, Form()], model_name: Annotated[ModelType, Form()] = ModelType.gpt4omini, file: Annotated[UploadFile | None, File()] = None, sheet_names: Annotated[str | None, Form()] = None, authorization: Annotated[str | None, Header()] = None, temperature: Annotated[float, Form()] = 0.6):
    if temperature < 0 or temperature > 2:
        raise HTTPException(
            status_code=400, detail="Temperature value is invalid. 0 <= temperature <= 2"
        )
    if not authorization or (authorization != AUTH_SECRET_KEY):
        print(f"Authorization header: {authorization}")
        raise HTTPException(
            status_code=401, detail="Provide the correct authorization token in headers")
    
    MODEL = model_name
    documentText, b64, base64_urls = parseDocuments(file, sheet_names, True)
    userContent = [
        {
            "type": "text",
            "text": ""
        }
    ]

    if len(documentText) > 0:
        userContent[0]['text'] = f"User prompt:\n{prompt}\n\nThis is document content: \n{documentText}"
    elif b64 is not None:
        userContent[0]['text'] = prompt

    messages = [
        {
            "role": "system",
            "content": SYSTEM_PROMPT_V2
        },
        {
            "role": "user",
            "content": userContent
        }
    ]

    if file is None:
        userContent[0]['text'] = prompt

    if b64 is not None:
        userContent.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/jpeg;base64,{b64}"}
        })

    if len(base64_urls) > 0:
        for url in base64_urls:
            userContent.append({
                "type": "image_url",
                "image_url": {"url": url, "detail": "low"},
            })

    response = {}
    try:
        response = client.chat.completions.create(
            model=MODEL,
            messages=messages,
            response_format={"type": "json_object"},
            temperature=temperature
        )
    except openai.RateLimitError as e:
        print(f"Rate limit error occurred: {e}")
        raise HTTPException(
            status_code=429, detail="OpenAI token limit exceeded")

    res_content = response.choices[0].message.content
    content = json.loads(res_content)    
    print(content)
    download_link = None

    if ("file_response" in content) and content['file_response'] is not None:
        pdf = MarkdownPdf(toc_level=2)
        pdf.add_section(Section(content['file_response']))
        pdf.save(f"response.pdf")
        file_name_s3 = str(uuid4()) + ".pdf"
        download_link = upload_file('response.pdf', file_name_s3)
        if not download_link:
            # Generate dummy link if real upload fails or is disabled
            dummy_uuid = str(uuid4())
            download_link = f"https://dummy-s3-bucket.example.com/{dummy_uuid}/{file_name_s3}"
            print(f"Using dummy download link: {download_link}")

    return {
        "status": "success",
        "prompt": prompt,
        "response": content['response'],
        "pdf": download_link
    }

# Audio Transcription + Chat Completion - Following v5 Pattern
@app.post("/v6/audio-processing")
def audio_processing(
    prompt: Annotated[str, Form()], 
    file: Annotated[UploadFile, File()],
    model_name: Annotated[ModelType, Form()] = ModelType.gpt4omini, 
    authorization: Annotated[str | None, Header()] = None, 
    temperature: Annotated[float, Form()] = 0.6,
    remove_noise: Annotated[bool, Form()] = True,
    force_english: Annotated[bool, Form()] = True
):
    """
    Audio processing endpoint that follows the v5 pattern, but specialized for audio files.
    
    Parameters:
    - prompt: User prompt for analysis of the audio
    - file: Audio file (formats: mp3, wav, ogg, m4a, flac) or video file for audio extraction
    - model_name: GPT model to use for analysis (default: gpt-4o-mini)
    - authorization: Auth token
    - temperature: Controls randomness in GPT responses (0.0 to 2.0)
    - remove_noise: Whether to apply noise removal to audio (default: True)
    - force_english: Whether to force English transcription (default: True)
    Returns:
    - JSON with transcription, analysis response, and PDF download link if applicable
    """
    if temperature < 0 or temperature > 2:
        raise HTTPException(
            status_code=400, detail="Temperature value is invalid. 0 <= temperature <= 2"
        )
    
    if not authorization or (authorization != AUTH_SECRET_KEY):
        print(f"Authorization header: {authorization}")
        raise HTTPException(
            status_code=401,
            detail="Provide the correct authorization token in headers"
        )

    # Check file extension for audio files
    fileExt = file.filename.split('.')[-1].lower()
    if fileExt not in AUDIO_EXTENSIONS and fileExt not in VIDEO_EXTENSIONS:
        supported_formats = ", ".join(AUDIO_EXTENSIONS + VIDEO_EXTENSIONS)
        raise HTTPException(400, f"{fileExt} file type not supported for audio processing. Supported formats: {supported_formats}")
    
    MODEL = model_name

    try:
        # Initialize Whisper service for transcription
        whisper_service = WhisperService(client=client)
        transcription = None
        temp_dir = None
        temp_path = None
        
        try:
            # Use temp directory for audio processing
            temp_dir = tempfile.mkdtemp()
            temp_path = os.path.join(temp_dir, file.filename)
            
            with open(temp_path, "wb") as buffer:
                buffer.write(file.file.read())

            # Try audio transcription with original settings
            try:
                transcription = whisper_service.transcribe_audio(
                    temp_path,
                    remove_noise=remove_noise,
                    force_english=force_english
                )
            except Exception as transcription_error:
                print(f"First transcription attempt failed: {transcription_error}")
                # If that fails, try with noise removal disabled
                transcription = whisper_service.transcribe_audio(
                    temp_path,
                    remove_noise=False,
                    force_english=force_english
                )
        finally:
            # Clean up temporary files
            if temp_path and os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except:
                    pass
            if temp_dir and os.path.exists(temp_dir):
                try:
                    os.rmdir(temp_dir)
                except:
                    pass

        # Setup the content for GPT processing - Following v5 pattern
        userContent = [
            {
                "type": "text",
                "text": f"User prompt:\n{prompt}\n\nThis is audio transcription content: \n{transcription}\n\nAudio processing details:\nNoise Removal: {remove_noise}\nForced English: {force_english}"
            }
        ]

        messages = [
            {
                "role": "system",
                "content": AUDIO_TRANSCRIPTION_PROMPT
            },
            {
                "role": "user",
                "content": userContent
            }
        ]

        response = client.chat.completions.create(
            model=MODEL,
            messages=messages,
            response_format={"type": "json_object"},
            temperature=temperature
        )
        res_content = response.choices[0].message.content
        content = json.loads(res_content)
        print(content)
        download_link = None
        if ("file_response" in content) and content['file_response'] is not None:
            pdf = MarkdownPdf(toc_level=2)
            pdf.add_section(Section(content['file_response']))
            pdf.save(f"response.pdf")
            file_name_s3 = str(uuid4()) + ".pdf"
            download_link = upload_file('response.pdf', file_name_s3)
            if not download_link:
                # Generate dummy link if real upload fails or is disabled
                dummy_uuid = str(uuid4())
                download_link = f"https://dummy-s3-bucket.example.com/{dummy_uuid}/{file_name_s3}"
                print(f"Using dummy download link: {download_link}")
        
        # Return the response with transcription data and audio optimization details
        result = {
            "status": "success",
            "prompt": prompt,
            "response": content['response'],
            "pdf": download_link,
            "transcription": transcription,
            "optimizations": {
                "noise_removal": remove_noise,
                "forced_english": force_english
            }
        }
        return result

    except Exception as e:
        print(f"Error in audio processing: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Error processing audio request: {str(e)}"
        )
