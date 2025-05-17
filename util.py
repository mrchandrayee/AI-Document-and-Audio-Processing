from fastapi import UploadFile, HTTPException
from pypdf import PdfReader
import tempfile
import base64
from striprtf.striprtf import rtf_to_text
from docx import Document
import pandas as pd
from io import BytesIO
import pypdfium2 as pdfium
import os
import boto3
from openai import OpenAI

SUPPORTED_EXTENSIONS = set(
    ['doc', 'dot', 'docx', 'dotx', 'docm', 'dotm', 'pdf', 'png', 'jpeg', 'jpg', 'rtf', 'xlsx', 'xls', 'txt',
     'mp3', 'wav', 'ogg', 'm4a', 'flac'])  # Added audio file extensions

def parseDocuments(file: UploadFile, sheet_names: str, parseAsImage: bool = False):
    documentText = ''
    base64_urls = []
    b64 = None

    if file is not None:
        fileExt = file.filename.split('.')[-1].lower()
        if fileExt not in SUPPORTED_EXTENSIONS:
            raise HTTPException(400, f"{fileExt} file type not supported")
        if fileExt == 'pdf':
            if not parseAsImage:
                reader = PdfReader(file.file)
                for page in reader.pages:
                    extracted_text = page.extract_text(extraction_mode='layout')
                    if len(extracted_text.strip()) == 0:
                        extracted_text = page.extract_text()
                    documentText += extracted_text
                    documentText += '\n\n'
                if len(documentText.strip()) == 0:
                    raise HTTPException(
                        status_code=400, detail="Unable to parse text from the given document or the document does not contain any text.")
            else:
                p = pdfium.PdfDocument(file.file)
                if len(p) > 20:
                    raise HTTPException(status_code=400, detail="File is too large to be processed")
                for i in range(len(p)):
                    page = p[i]
                    buffered = BytesIO()
                    image = page.render(scale=2).to_pil()
                    image.save(buffered, format="JPEG")
                    img_str = base64.b64encode(buffered.getvalue()).decode("utf-8")
                    base64_urls.append(f"data:image/jpeg;base64,{img_str}")
                documentText = "Images are provided as document"

        elif fileExt in ['png', 'jpeg', 'jpg']:
            with tempfile.TemporaryDirectory() as temp_dir:
                path = f"{temp_dir}/filename.{fileExt}"
                with open(path, 'wb') as img:
                    img.write(file.file.read())
                with open(path, 'rb') as img:
                    b64 = base64.b64encode(img.read()).decode("utf-8")
        elif fileExt in ['rtf']:
            documentText = rtf_to_text(file.file.read().decode('utf-8'))
        elif fileExt in ['xls', 'xlsx']:
            try:
                excelDF = None
                if sheet_names is None or len(sheet_names) == 0:
                    excelDF = pd.read_excel(file.file)
                    documentText = excelDF.to_string()
                else:
                    sheet_name_arr = list(map(lambda x: x.strip(), sheet_names.split(',')))
                    excelDF = pd.read_excel(file.file, sheet_name=sheet_name_arr)
                    if not isinstance(excelDF, dict):
                        raise HTTPException(400, "Something went wrong while parsing the excel sheet")
                    for key in excelDF:
                        documentText += f'Sheet Name: {key}\n\n'
                        documentText += excelDF[key].to_string()
                        documentText += '\n\n'
            except Exception as exp:
                print(exp)
                raise HTTPException(400, "Worksheet name not found")
        elif fileExt == 'txt':
            documentText += file.file.read().decode('utf-8')
        else:
            try:
                document = Document(file.file)
            except Exception as e:
                print("Something went wrong while parsing the file:", e)
                raise HTTPException(400, detail="The file could not be parsed")
            for paragraph in document.paragraphs:
                documentText += paragraph.text if paragraph and paragraph.text else ""
            # text = ''
            # for elem in document.tables:
            #     textRow = set()
            #     for i, row in enumerate(elem.columns):
            #         for cell in row.cells:
            #             textRow.add(f"{cell.text}\n")
            #     text += f"{' '.join(list(textRow))}\n\n"
            # print(text)
    return [documentText, b64, base64_urls]

def parseDocumentsV2(file: UploadFile, sheet_names: str, client: OpenAI = None):
    documentText = ''
    b64 = None
    pdf_file_id = None

    if file is not None:
        fileExt = file.filename.split('.')[-1].lower()
        if fileExt not in SUPPORTED_EXTENSIONS:
            raise HTTPException(400, f"{fileExt} file type not supported")
        if fileExt == 'pdf':
            if client is not None:
                path = f"upload_filename.{fileExt}"
                print("Started uploading....")
                with open(path, 'wb') as f:
                    f.write(file.file.read())
                with open(path, 'rb') as f:
                    uploaded_file = client.files.create(
                        file = f,
                        purpose="user_data"
                    )
                    print("Compeleted uploading....", uploaded_file.id)
                    pdf_file_id = uploaded_file.id
                os.remove(path)
        elif fileExt in ['png', 'jpeg', 'jpg']:
            with tempfile.TemporaryDirectory() as temp_dir:
                path = f"{temp_dir}/filename.{fileExt}"
                with open(path, 'wb') as img:
                    img.write(file.file.read())
                with open(path, 'rb') as img:
                    b64 = base64.b64encode(img.read()).decode("utf-8")
        elif fileExt in ['rtf']:
            documentText = rtf_to_text(file.file.read().decode('utf-8'))
        elif fileExt in ['xls', 'xlsx']:
            try:
                excelDF = None
                if sheet_names is None or len(sheet_names) == 0:
                    excelDF = pd.read_excel(file.file)
                    documentText = excelDF.to_string()
                else:
                    sheet_name_arr = list(map(lambda x: x.strip(), sheet_names.split(',')))
                    excelDF = pd.read_excel(file.file, sheet_name=sheet_name_arr)
                    if not isinstance(excelDF, dict):
                        raise HTTPException(400, "Something went wrong while parsing the excel sheet")
                    for key in excelDF:
                        documentText += f'Sheet Name: {key}\n\n'
                        documentText += excelDF[key].to_string()
                        documentText += '\n\n'
            except Exception as exp:
                print(exp)
                raise HTTPException(400, "Worksheet name not found")
        elif fileExt == 'txt':
            documentText += file.file.read().decode('utf-8')
        else:
            try:
                document = Document(file.file)
            except Exception as e:
                print("Something went wrong while parsing the file:", e)
                raise HTTPException(400, detail="The file could not be parsed")
            for paragraph in document.paragraphs:
                documentText += paragraph.text if paragraph and paragraph.text else ""
    return [documentText, b64, pdf_file_id]

def parseDocumentsWithVector(file: UploadFile, sheet_names: str, client: OpenAI = None):
    documentText = ''
    b64 = None
    vector_store_id = None

    if file is not None:
        fileExt = file.filename.split('.')[-1].lower()
        if fileExt not in SUPPORTED_EXTENSIONS:
            raise HTTPException(400, f"{fileExt} file type not supported")
        if fileExt == 'pdf':
            vector_store = client.vector_stores.create(
                name="API Vector Store",
                expires_after={
                    'anchor': "last_active_at",
                    'days': 1
                }
            )
            vector_store_id = vector_store.id
            if client is not None:
                path = f"upload_filename.{fileExt}"
                print("Started uploading....")
                with open(path, 'wb') as f:
                    f.write(file.file.read())
                with open(path, 'rb') as f:
                    uploaded_file = client.files.create(
                        file = f,
                        purpose="user_data"
                    )
                    print("Compeleted uploading....", uploaded_file.id)
                    pdf_file_id = uploaded_file.id

                client.vector_stores.files.create(
                    vector_store_id=vector_store_id,
                    file_id=pdf_file_id
                )
                os.remove(path)
        elif fileExt in ['png', 'jpeg', 'jpg']:
            with tempfile.TemporaryDirectory() as temp_dir:
                path = f"{temp_dir}/filename.{fileExt}"
                with open(path, 'wb') as img:
                    img.write(file.file.read())
                with open(path, 'rb') as img:
                    b64 = base64.b64encode(img.read()).decode("utf-8")
        elif fileExt in ['rtf']:
            documentText = rtf_to_text(file.file.read().decode('utf-8'))
        elif fileExt in ['xls', 'xlsx']:
            try:
                excelDF = None
                if sheet_names is None or len(sheet_names) == 0:
                    excelDF = pd.read_excel(file.file)
                    documentText = excelDF.to_string()
                else:
                    sheet_name_arr = list(map(lambda x: x.strip(), sheet_names.split(',')))
                    excelDF = pd.read_excel(file.file, sheet_name=sheet_name_arr)
                    if not isinstance(excelDF, dict):
                        raise HTTPException(400, "Something went wrong while parsing the excel sheet")
                    for key in excelDF:
                        documentText += f'Sheet Name: {key}\n\n'
                        documentText += excelDF[key].to_string()
                        documentText += '\n\n'
            except Exception as exp:
                print(exp)
                raise HTTPException(400, "Worksheet name not found")
        elif fileExt == 'txt':
            documentText += file.file.read().decode('utf-8')
        else:
            try:
                document = Document(file.file)
            except Exception as e:
                print("Something went wrong while parsing the file:", e)
                raise HTTPException(400, detail="The file could not be parsed")
            for paragraph in document.paragraphs:
                documentText += paragraph.text if paragraph and paragraph.text else ""
    return [documentText, b64, vector_store_id]

BUCKET_NAME = os.getenv("AWS_BUCKET_NAME")
AWS_REGION = os.getenv("AWS_BUCKET_REGION")

def upload_file(file_path, object_name=None):
    if object_name is None:
        object_name = os.path.basename(file_path)
    response = False
    s3_client = boto3.client('s3', region_name=AWS_REGION)
    try:
        s3_client.upload_file(file_path, BUCKET_NAME, object_name)
        response = f"https://{BUCKET_NAME}.s3.amazonaws.com/{object_name}"
    except Exception as e:
        print("Error while upload:", e)
        return False
    return response